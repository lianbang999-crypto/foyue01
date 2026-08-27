from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/bincai/Downloads/foyue/looka")
IMG_ROOT = ROOT / "参考组件图标"
OUT = ROOT / "docs" / "Lifebear_UIUX_逐页面图鉴与功能链路_2026-08-26.docx"

# compact_reference_guide preset + named overrides:
# - CJK glyph fallback: PingFang SC (ASCII remains Calibri)
# - Screen annotation matrix: 2.40in evidence + 4.10in structured observation
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
LINE = "D9DEE7"
WHITE = "FFFFFF"
BLACK = "111827"
CORAL = "D85D5D"
POSITIVE = "1F3A5F"
CAUTION = "7A5A00"
RISK = "9B1C1C"
CJK_FONT = "Noto Sans CJK SC"


def find_image(token: str) -> Path:
    hits = sorted(IMG_ROOT.rglob(f"*{token}*"))
    hits = [p for p in hits if p.suffix.lower() in {".jpg", ".jpeg", ".png"}]
    if len(hits) != 1:
        raise RuntimeError(f"Image token {token!r} resolved to {len(hits)} files: {hits}")
    return hits[0]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Iterable[int], indent=120, border=True) -> None:
    widths = list(widths)
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single" if border else "nil")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), LINE)


def apply_cjk_font(r_pr, font_name=CJK_FONT) -> None:
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def set_run_font(run, size=None, bold=None, color=None, italic=None, ascii_font=CJK_FONT):
    run.font.name = ascii_font
    apply_cjk_font(run._element.get_or_add_rPr(), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CJK_FONT
    apply_cjk_font(normal._element.get_or_add_rPr())
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = CJK_FONT
        apply_cjk_font(style._element.get_or_add_rPr())
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Screen Note" not in [s.name for s in styles]:
        style = styles.add_style("Screen Note", 1)
    else:
        style = styles["Screen Note"]
    style.font.name = CJK_FONT
    apply_cjk_font(style._element.get_or_add_rPr())
    style.font.size = Pt(9.2)
    style.font.color.rgb = RGBColor.from_string(BLACK)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.12

    if "Evidence Caption" not in [s.name for s in styles]:
        style = styles.add_style("Evidence Caption", 1)
    else:
        style = styles["Evidence Caption"]
    style.font.name = CJK_FONT
    apply_cjk_font(style._element.get_or_add_rPr())
    style.font.size = Pt(8)
    style.font.color.rgb = RGBColor.from_string(MUTED)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.0


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=8.5, color=MUTED)


def setup_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run("LIFEBEAR UI/UX 逐页面图鉴")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        add_page_field(fp)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_label_paragraph(container, label: str, text: str, *, color=INK) -> None:
    p = container.add_paragraph(style="Screen Note")
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=9.2, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=9.2, color=BLACK)


def add_callout(doc: Document, label: str, text: str, fill=CALLOUT, color=INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120, border=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=10, color=BLACK)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LIFEBEAR")
    set_run_font(r, size=12, bold=True, color=CORAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("UI/UX 逐页面图鉴\n与功能链路说明")
    set_run_font(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("基于现有 Lifebear 截图与录屏的页面级设计研究")
    set_run_font(r, size=13, color=MUTED)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [3120, 3120, 3120], indent=120, border=False)
    for idx, token in enumerate(("20260820131547_4_18", "20260820135053_8_4", "20260820133743_16_18")):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(find_image(token)), width=Inches(1.45))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("研究日期  2026-08-26")
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("用途：Looka 全局 UI/UX 统一设计、评审与验收")
    set_run_font(r, size=9.5, color=MUTED)


def add_front_matter(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("如何阅读这份图鉴", level=1)
    add_callout(doc, "证据标记", "B = 截图或录屏可确认；C = 从 Lifebear 归纳出的 Looka 统一建议。页面中的尺寸若写“约”，代表按 1136px 宽截图折算，而非官方设计稿标注。")
    doc.add_heading("覆盖范围", level=2)
    for label, text in (
        ("日历", "月 / 周 / 日视图、选日议程、视图切换、贴纸托盘与日期落点。"),
        ("日程", "新建、编辑、字段子编辑器、模板、提醒与删除确认。"),
        ("ToDo", "首页、清单、未来 7 天、星标、完成任务、任务编辑、排序、菜单与建清单。"),
        ("笔记与日记", "容器列表、内容列表、编辑器、搜索、排序、重命名与新建列表。"),
        ("更多与设置", "设置路由、账户、方案、主题/贴纸商店、帮助、消息与口令。"),
    ):
        p = doc.add_paragraph()
        r = p.add_run(f"{label}  ")
        set_run_font(r, bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(text), color=BLACK)
    doc.add_heading("全局视觉结论", level=2)
    add_callout(doc, "设计 DNA", "Calm Utility × Emotional Warmth。白底、轻分隔、低圆角、黑灰动作建立工具骨架；彩色分类、周末、贴纸与主题只在需要表达语义或情绪时出现。", fill=BLUE_GRAY)

    doc.add_page_break()
    doc.add_heading("全局功能架构与主导航", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1700, 1850, 3310, 2500])
    headers = ("入口", "交互类型", "主链路", "关键约束")
    for i, h in enumerate(headers):
        set_cell_shading(table.cell(0, i), BLUE_GRAY)
        p = table.cell(0, i).paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Calendar", "持续页签", "月/周/日 → 选日 → 议程/贴纸", "切日期不离开日历；保持上下文"),
        ("ToDo", "持续页签", "首页 → 清单/智能视图 → 任务", "完成/星标即时；编辑/删除进入动作层"),
        ("中央 +", "全局创建", "+ → 选择创建对象 → 编辑器", "高频入口居中；不承载二级设置"),
        ("Note & Diary", "持续页签", "容器 → 内容列表 → 编辑器", "列表持久化；搜索绑定当前域"),
        ("More", "路由页", "账户/方案/设置/商店/帮助", "低频功能集中；不可挤占高频底栏"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(text)
            set_run_font(r, size=9, color=BLACK)
    doc.add_paragraph()
    add_callout(doc, "全局返回规则", "X = 取消创建并返回来源页；← = 返回上一级且尽量保留筛选/滚动/选中日；系统返回键与标题栏返回行为必须一致。")

    doc.add_page_break()
    doc.add_heading("全局尺寸、圆角与字体基线", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2500, 2200, 4660])
    for i, h in enumerate(("对象", "建议冻结值", "依据与说明")):
        set_cell_shading(table.cell(0, i), BLUE_GRAY)
        r = table.cell(0, i).paragraphs[0].add_run(h)
        set_run_font(r, size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    token_rows = [
        ("屏幕左右安全边距", "16dp", "B/C：列表、编辑器与底栏对齐的共同起点。"),
        ("基础间距", "4 / 8 / 12 / 16 / 24dp", "C：所有垂直节奏只从该组组合。"),
        ("点击热区", "≥44×44dp", "C：图标可小，但热区不能小。"),
        ("表单行高", "48–56dp", "B/C：单行选择、开关与摘要行。"),
        ("顶栏", "44–56dp", "B/C：标题、关闭、保存保持同一基线。"),
        ("底部导航", "约58dp", "B：五入口；中央 + 视觉直径约48dp。"),
        ("标准 Dialog 宽度", "约屏宽 86%", "B：1136px 证据中约977px；两侧约26dp。"),
        ("Dialog 圆角", "约2dp", "B：实机为低圆角，不是 16–28dp 大圆角卡片。"),
        ("Dialog 内边距", "约24–26dp", "B/C：标题、内容、动作稳定对齐。"),
        ("标题层级", "32 / 20 / 18sp", "B/C：月份 / 弹窗标题 / 页面或空态标题。"),
        ("正文与行标签", "16 / 14–16sp", "B/C：正文与高频行；弱信息 12sp。"),
        ("日历日期与日程", "10–12 / 8–10sp", "B/C：密度高，但需确保可扫读。"),
    ]
    for row in token_rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            r = cells[i].paragraphs[0].add_run(text)
            set_run_font(r, size=8.9, color=BLACK)

    doc.add_page_break()
    doc.add_heading("弹窗、Bottom Sheet 与反馈层级", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1250, 1900, 3110, 3100])
    for i, h in enumerate(("层级", "组件", "适用", "交互规则")):
        set_cell_shading(table.cell(0, i), BLUE_GRAY)
        set_run_font(table.cell(0, i).paragraphs[0].add_run(h), size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("L0", "行内反馈", "开关、完成、星标", "立即更新；失败原位回滚。"),
        ("L1", "Snackbar", "成功、撤销", "不夺焦点；约160–200ms 进入。"),
        ("L2", "Banner", "账户/订阅/能力提示", "可持续，但不可遮住主任务。"),
        ("L3", "Picker / Sheet", "视图、清单、日期等选择", "一个临时选择上下文；保留背景位置。"),
        ("L4", "Dialog", "确认、单/多选、危险动作", "遮罩约60%黑；一次只出现一个。"),
        ("L5", "系统层", "权限、系统键盘", "应用层主动提示必须退让。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_run_font(cells[i].paragraphs[0].add_run(text), size=8.9, color=BLACK)
    doc.add_paragraph()
    add_callout(doc, "高度策略", "单选/删除确认约18%屏高；范围选择约20%；视觉选择约45%；展开颜色约57%；视图切换 Bottom Sheet 约41%。高度由内容类型决定，不做统一固定高度。")


SCREENS = [
    dict(section="A 日历", title="A1 月视图：全局时间骨架", image="20260820131547_4_18", role="月级浏览与跨功能回落首页。核心不是“展示更多”，而是用最低噪声让日期、周末、节日和少量事项同时可扫读。", layout="月份标题位于左上；7 列网格占主区；日期数字与事项色条分层；底栏固定。大面积白底保留呼吸感，不使用卡片包裹日历。", detail="月份约28–32sp；日期约10–12sp；事项约8–10sp；周末/节日用受控红蓝；分隔线极浅。选中态以细描边或浅底表达。", keys="日历图标=当前页；日期单击=选中并更新下方/后续议程；右上日历/今天=回到今天；中央 +=进入创建对象选择。", chain="应用启动/底栏 Calendar → 月视图 → 点日期 → 保持月视图上下文并刷新该日议程。", state="选中日需要唯一；跨月滑动后“今天”入口仍可见；空日期不弹空白 Dialog。", looka="冻结月视图为白底网格，不将每个日期做成圆角卡片。"),
    dict(section="A 日历", title="A2 选日 + 议程：主视图不丢上下文", image="20260820131551_5_18", role="把“看月份”和“处理某一天”合成同一持续页面，减少进入详情再返回的成本。", layout="上部月网格，下部选中日议程；中间用细分隔/拖柄建立主从关系。议程为空时仍保留日期标题和日记提示。", detail="选中日期使用低对比描边；议程行采用文字+少量色点，不引入厚阴影；空态文案比月份标题小一个层级。", keys="日期=切换议程；事项行=打开详情；空白区域/+=创建；底栏切页后返回应恢复原选中日。", chain="月视图 → 点日期 → 议程区更新 → 点事项 → 详情 → 返回 → 原日期与滚动位置。", state="日期切换应即时；加载时只在议程区显示骨架，不锁住整页。", looka="采用“上时间、下行动”的持续上下文模型。"),
    dict(section="A 日历", title="A3 周视图：高密度时间轴", image="20260820133739_13_18", role="面向一周内的冲突判断与时间分配。", layout="顶部 7 日列；左侧小时轴；日程块贴合时间位置。白底+细网格，颜色只用于事件类别和冲突识别。", detail="小时文字约10sp；事件色块小圆角或近直角；同列重叠时缩窄并排，不用浮层覆盖整行。", keys="日期头=聚焦某日；事件块=详情；空白时间槽=以该时间预填创建；右上视图键=打开视图切换 Sheet。", chain="月视图 → 视图切换 → 周 → 点事件/时间槽 → 详情或已预填编辑器。", state="当前时间用细红线或点；跨周后今天按钮保留。", looka="时间轴信息密度高，避免装饰背景和大圆角。"),
    dict(section="A 日历", title="A4 日视图：单日纵向安排", image="20260820133741_15_18", role="处理单日精细时间安排，优先读时段和冲突。", layout="单列时间轴最大化；顶部仍保留日期上下文；底栏不变。", detail="垂直节奏由小时线主导；无事件区域保持纯白；返回月/周不改变底栏位置。", keys="日期标题/左右切换=前后一天；时间槽=新建；事件=详情；视图键=月/周/日切换。", chain="周/月 → 视图 Sheet → 日 → 滑动日期 → 点时间槽 → 日程创建（日期时间已带入）。", state="跨天切换不应闪回月视图；当前时间位置可见。", looka="日视图与周视图共享时间轴 Token。"),
    dict(section="A 日历", title="A5 视图切换 Bottom Sheet", image="20260820133740_14_18", role="在保留日历背景与当前选择的前提下切换月/周/日，以及进入日期移动、显示设置。", layout="底部 Sheet 全宽，约屏高40.7%；背景整体压暗。列表行近直角，当前项使用整行浅灰选中。", detail="证据中 Sheet 顶部视觉接近直边；不要默认套 20dp 大圆角。行高约48–56dp，图标与文字左对齐。", keys="月/周/日=选择即关闭并切换；日期移动=进入日期定位；显示设置=进入全页设置；点遮罩/返回=关闭不改变视图。", chain="日历右上视图键 → Sheet → 选择周 → Sheet 关闭 → 周视图，原日期保留。", state="只允许一个 Sheet；打开时底栏不可再触发新 Overlay。", looka="Sheet 用于小规模导航选择，不承载复杂表单。"),
    dict(section="B 日程", title="B1 新建日程：父编辑器", image="20260822175134_14_13", role="集中编辑日程的交易型父页面；子字段以摘要行进入各自编辑器。", layout="顶栏 X / 页面标题 / 保存；标题输入优先自动聚焦；日期、分类、提醒、重复、地点、备注依次纵排。", detail="保存按钮是唯一主动作；未填必填项时弱化。每行48–56dp，图标低对比，摘要靠近字段名。", keys="X=取消并回来源；保存=一次性提交父表单；相机=附图/能力门；日期/提醒/重复等行=进入子编辑器。", chain="中央 + → 日程 → 新建页 → 子编辑器修改字段 → 子编辑器完成 → 回父页 → 保存 → 回来源日历。", state="子编辑器“完成”只提交单字段；父页“保存”才提交整条日程；脏数据退出需确认。", looka="严格区分字段完成与整单保存。"),
    dict(section="B 日程", title="B2 日程创建 + 键盘：焦点所有权", image="20260823185620_56_13", role="展示标题输入时的键盘、底部快捷类型与表单压缩方式。", layout="键盘出现后内容区向上调整；标题光标可见；保存始终留在顶栏，不被键盘遮挡。", detail="IME 是唯一焦点拥有者；弹窗不可在键盘上再叠第二个输入层。", keys="键盘完成=结束当前字段输入，不保存整条；X=退出；保存=提交；底部日历/任务/贴纸图标=切换创建类型。", chain="创建入口 → 标题自动聚焦 → 键盘输入 → 点字段关闭/迁移焦点 → 保存。", state="切创建类型前需保留或确认当前草稿；系统返回先收键盘，再处理离页。", looka="统一键盘出现时的内容避让和返回优先级。"),
    dict(section="B 日程", title="B3 详情 / 编辑：查看与修改共用模型", image="20260822202411_24_4", role="任务/日程详情强调对象状态与低频动作，不让编辑控件长期占据阅读界面。", layout="顶栏返回、标题、动作图标；主体以状态、日期、所属清单/分类、提醒等字段组成。", detail="查看态文字比编辑态更安静；编辑、复制、删除集中在顶栏或更多菜单。", keys="←=回来源；编辑=进入编辑态；更多=复制/删除；星标/完成=即时状态动作。", chain="列表/日历 → 详情 → 编辑 → 保存 → 详情或来源；删除 → 确认 → 来源 + Snackbar。", state="即时状态失败应原位回滚；删除后来源列表数量同步。", looka="阅读态与编辑态视觉密度分开。"),
    dict(section="B 日程", title="B4 删除确认：危险动作的最后闸门", image="20260823182237_52_13", role="确认不可逆或高损失的日程删除。", layout="屏宽约86%的标准 Dialog，约18%屏高；遮罩约60%黑；标题、说明与动作按24–26dp内边距对齐。", detail="低圆角约2dp；危险动作使用明确文字，不依赖红色图标猜测；取消与删除分离。", keys="取消/点外部=关闭且不变；删除=提交删除；系统返回=等同取消。", chain="日程详情 → 更多/删除 → 确认 Dialog → 删除 → 返回日历 + 可选撤销提示。", state="确认层出现后背景不可交互；删除中主动作防重复点击。", looka="所有删除确认统一宽度、遮罩、动作顺序和文案结构。"),
    dict(section="C ToDo", title="C1 ToDo 首页：智能视图 + 用户清单", image="20260820135053_8_4", role="把高频智能视图、用户清单、标签与完成归档放在一个可扫读入口。", layout="顶部搜索；智能视图在前，用户清单居中，完成项在后；色点/星标/日历图标承担语义。", detail="白底列表、轻分隔；数量靠右；“创建”是文字行而非大卡片。", keys="搜索=当前任务域检索；智能视图=过滤；清单=进入；创建清单=Dialog；完成任务/完成清单=归档页。", chain="底栏 ToDo → 首页 → 选智能视图/清单 → 任务列表 → 详情/编辑。", state="数量为0仍可进入；能力受限时在动作发生点弹能力门。", looka="信息架构分组靠留白和小标题，不靠卡片。"),
    dict(section="C ToDo", title="C2 清单详情：完成、提醒、星标分工", image="20260824171113_63_4", role="在一个清单内快速扫读和维护任务。", layout="顶栏返回、清单名、更多；首行“添加任务”；任务行左完成、右提醒/星标。", detail="完成圆点是主要状态；星标只占右侧固定位置；拖拽把手只在排序态出现。", keys="完成=即时切换；星标=即时切换；任务行=详情；添加任务=新建并预选当前清单；更多=编辑/排序。", chain="ToDo 首页 → 清单 → 任务行 → 详情；或添加任务 → 新建 → 保存 → 当前清单。", state="完成后可移入完成区或淡出；失败原位回滚。", looka="高频状态动作就地完成，低频管理进菜单。"),
    dict(section="C ToDo", title="C3 未来 7 天：时间型智能视图", image="20260824171108_58_4", role="跨清单查看即将到期任务；日期分组优先于清单归属。", layout="标题 + 按日期分组的任务行；所属清单以弱文字/色点呈现。", detail="不同清单不做彩色卡片；星标仍在固定右侧，完成仍在固定左侧。", keys="任务行=详情；完成/星标=即时；+=以智能视图上下文创建。", chain="ToDo 首页 → 未来7天 → 任务 → 编辑 → 返回仍在未来7天。", state="过期任务需有清晰标记，但不与错误状态混色。", looka="智能视图只改变排序/筛选，不改变任务行组件。"),
    dict(section="C ToDo", title="C4 星标：优先级视图与空态", image="20260824171109_59_4", role="聚合所有星标任务；证据画面同时展示空态处理。", layout="顶部快速添加与星标；空态居中插图+短文案，底栏和顶栏不变。", detail="空态插图低饱和；不使用全屏弹窗教育；空态标题约18sp。", keys="快速添加=创建任务并默认星标；更多=编辑；底栏返回其他域。", chain="ToDo 首页 → 星标 → 空态/列表 → 添加任务 → 保存 → 星标列表。", state="空态、加载态、错误态必须区分；空态可直接行动。", looka="空态给下一步，不堆功能说明。"),
    dict(section="C ToDo", title="C5 完成任务：归档阅读", image="20260824174422_81_4", role="按时间范围回看已完成任务，不混入当前执行列表。", layout="顶部范围选择；下方按日期分组；完成图标为稳定已完成态。", detail="已完成项仍保持可读，避免过度降低透明度；提醒与星标保持弱显示。", keys="范围=打开范围 Dialog；任务行=详情；返回=ToDo 首页。", chain="ToDo 首页 → 完成任务 → 范围选择 → 选择后刷新 → 点任务查看。", state="切范围时保留滚动策略要一致；空范围展示归档空态。", looka="完成归档与当前任务视觉分区。"),
    dict(section="C ToDo", title="C6 完成范围选择：单选即提交", image="20260824174421_80_4", role="选择完成任务的时间窗口。", layout="标准宽度 Dialog，约20%屏高；选项整行，当前项整行浅灰。", detail="单选结果可“选择即关闭”；不额外放冗余 OK。", keys="选项=立即提交并关闭；点外部/返回=取消不变。", chain="完成任务页 → 点范围 → Dialog → 选“过去3个月/某年” → 关闭并刷新。", state="打开时当前值必须高亮；关闭后焦点回范围控件。", looka="单选路由采用 select-and-close。"),
    dict(section="C ToDo", title="C7 更多菜单：低频动作锚定", image="20260824171110_60_4", role="清单页将编辑等低频动作放入锚定菜单，避免常驻抢占。", layout="菜单锚定右上更多键；宽度只容纳动作文字；背景不做重遮罩。", detail="动作数量少、顺序稳定；危险项如删除需二次确认。", keys="编辑=进入清单编辑；点外部/返回=关闭；其余动作按证据路由。", chain="清单 → 更多 → 编辑/排序/删除 → 目标页或确认层 → 返回清单。", state="菜单打开时第二次点更多关闭；不与 Dialog 同时存在。", looka="锚定菜单只放低频、对象级动作。"),
    dict(section="C ToDo", title="C8 排序：进入明确模式", image="20260824171111_61_4", role="任务重排是一个暂时模式，不与普通浏览混淆。", layout="标题改为排序；每行右侧出现拖拽把手；其他低频按钮弱化。", detail="拖动反馈依赖位移与占位线，不用大阴影卡片。", keys="拖拽把手=重排；返回=提交或取消须固定一种策略；任务内容在排序态不进入详情。", chain="清单 → 更多 → 排序 → 拖动 → 返回/完成 → 清单新顺序。", state="排序过程中完成/星标禁用，避免状态与位置同时改变。", looka="显式模式切换，避免手势歧义。"),
    dict(section="C ToDo", title="C9 建清单：完整颜色面板", image="20260826115525_96_4", role="为新清单命名并选择语义色。", layout="约86%宽 Dialog；输入在上、颜色网格在中、动作为底；展开色板约57%屏高。", detail="颜色圆点间距均匀；当前色用环/勾双重表达；低圆角、强遮罩。", keys="名称=自动聚焦；色彩=即时预览；取消=不创建；保存=创建并返回/选中。", chain="ToDo 首页或清单选择器 → 创建清单 → 输入/选色 → 保存 → 新清单持久化并自动可用。", state="空名称时保存禁用；键盘和色板不能抢占同一高度。", looka="颜色只做分类语义，不做整屏染色。"),
    dict(section="C ToDo", title="C10 建清单 + 键盘：紧凑态", image="20260826115526_97_4", role="输入阶段收起完整色板，优先保证键盘与命名任务。", layout="键盘出现后 Dialog 约27.5%屏高；仅显示当前色与展开箭头。", detail="Dialog 上移但不贴顶；动作始终可见；返回键先收键盘。", keys="颜色箭头=收起键盘/展开色板；保存=创建；取消=丢弃。", chain="创建清单 → 名称聚焦 → 输入 → 选色/保存 → 新清单。", state="IME、色板、父 Dialog 三者只有一个焦点所有者。", looka="输入态与视觉选择态用高度模式切换。"),
    dict(section="C ToDo", title="C11 能力门：在触发点解释限制", image="20260820135051_6_4", role="用户触发高级标签/功能时，解释限制并提供关闭或了解方案。", layout="标准 Dialog + 插图/说明 + 双动作；背景保持任务上下文。", detail="能力门不是广告首页；主动作明确“详细了解/注册”等后续。", keys="关闭=返回原动作点；了解/注册=进入方案/账户；系统返回=关闭。", chain="受限动作 → 能力门 → 关闭或方案/账户 → 返回原模块。", state="同一会话避免重复打断；成功解锁后应恢复原意图。", looka="能力提示只在用户表达意图后出现。"),
    dict(section="D 笔记与日记", title="D1 笔记容器页：列表是持久对象", image="20260820133743_16_18", role="先选择笔记列表，再进入内容；列表不是一次性筛选项。", layout="顶部搜索；笔记/日记双页签；列表行带容器图标；“创建列表”作为末尾文字动作。", detail="选中页签使用细下划线；列表数量弱化靠右；无卡片边框。", keys="搜索=标题/正文；页签=笔记/日记切换；列表=进入内容；创建列表=Dialog。", chain="底栏 Note&Diary → 笔记页签 → 列表 → 笔记内容 → 编辑器。", state="新列表保存后持久出现并自动选中；返回恢复原列表。", looka="容器与筛选概念分开。"),
    dict(section="D 笔记与日记", title="D2 日记索引：日期优先", image="20260820133743_17_18", role="日记按月份和日期组织，强调时间回看而非文件夹。", layout="搜索与双页签不变；内容按月份分段，日期数字形成视觉锚点。", detail="月份标签小而弱；日期数字18sp左右；标题与星期更小。", keys="搜索=日记正文；条目=详情；页签=返回笔记；+=新建当天/指定日记。", chain="Note&Diary → 日记 → 条目 → 详情/编辑 → 返回原月份。", state="同日多条的排序规则需稳定；空月份不展示空分组。", looka="日记采用时间索引，不复制笔记文件夹结构。"),
    dict(section="D 笔记与日记", title="D3 笔记列表详情：对象级菜单", image="20260824164354_78_13", role="显示当前列表内的笔记，并提供添加、排序、重命名等对象级管理。", layout="标题栏为列表名；内容行极简；右上更多承载管理动作。", detail="列表为空时仍保留添加入口；广告/系统横幅属于L2层，不应改变内容结构。", keys="笔记行=详情；+=新建；更多=编辑/排序/重命名；返回=容器页。", chain="笔记容器 → 列表 → 笔记 → 返回列表；更多 → 管理动作 → 返回列表。", state="排序/重命名完成后列表名和顺序即时同步。", looka="对象级管理不常驻在每一行。"),
    dict(section="D 笔记与日记", title="D4 笔记详情 / 编辑器", image="20260824164355_79_13", role="轻量纯文本记录；阅读优先，必要时进入编辑。", layout="返回、标题/日期、相机、编辑、更多位于顶栏；正文占满白底。", detail="正文16sp左右、行距宽松；工具按钮不加底色；选择文本时才出现上下文动作。", keys="返回=列表；相机=附图；编辑=进入输入态；更多=删除/移动等；保存=提交。", chain="列表 → 笔记详情 → 编辑 → 保存 → 详情/列表。", state="脏内容退出需确认；键盘返回先收起，再离页。", looka="编辑器避免多层工具栏，保留纯净书写面。"),
    dict(section="D 笔记与日记", title="D5 日记编辑器：日期即主键", image="20260824164359_83_13", role="围绕日期记录内容，允许相机与编辑动作，但不增加复杂模板壳。", layout="标题栏显示日期；正文留白大；相机、编辑、更多保持右侧固定。", detail="日期比正文标题更重要；日记条目之间依赖日期索引，不依赖彩色卡片。", keys="日期标题=可切换/选择日期；相机=媒体；编辑=输入；更多=对象动作。", chain="日记索引 → 某日 → 编辑 → 保存 → 返回该月位置。", state="修改日期时避免复制出重复日记；冲突需明确提示。", looka="日记与笔记共享编辑器骨架，仅信息架构不同。"),
    dict(section="D 笔记与日记", title="D6 排序 Dialog：多选式提交动作", image="20260824164352_76_13", role="选择列表内容的排序依据。", layout="标准宽度 Dialog；标题、单选项、底部取消/应用。", detail="当前项用 radio + 文案；因为需要显式应用，改变先留在草稿态。", keys="单选=更新草稿；取消=丢弃；应用=提交排序并关闭。", chain="列表更多 → 排序 → 选择依据 → 应用 → 列表重排。", state="重开 Dialog 应显示已提交值，而非上次取消的草稿。", looka="需要预览/多步的选择使用 draft + OK。"),
    dict(section="D 笔记与日记", title="D7 重命名 Dialog：短输入事务", image="20260824164353_77_13", role="修改持久列表名称。", layout="标题、单行输入、取消/重命名；键盘出现后仍保持动作可见。", detail="约86%屏宽；输入下划线比外框更轻；错误在输入行下方。", keys="取消=不变；重命名=校验并提交；回车可等同主动作但不能绕过校验。", chain="列表更多 → 重命名 → 输入 → 重命名 → 标题与容器页同步。", state="空值/重复名/超长即时提示；提交中防重复。", looka="短事务用 Dialog，不跳全页。"),
    dict(section="D 笔记与日记", title="D8 新建列表 Dialog：创建后自动选中", image="20260824164356_80_13", role="在笔记域中新建持久列表，并把它作为当前选择。", layout="紧凑 Dialog + 单行输入 + 键盘；约27.5%屏高。", detail="输入优先；动作区与键盘之间留安全距；低圆角。", keys="取消=丢弃；保存=创建；系统返回先收键盘，再关闭。", chain="笔记容器/列表选择 → 创建列表 → 输入 → 保存 → 新列表出现并自动进入/选中。", state="创建成功后不能只更新临时选择；必须写入容器列表。", looka="把“建容器”和“换筛选”从数据模型上分开。"),
    dict(section="E 贴纸", title="E1 贴纸托盘：情绪层不破坏工具骨架", image="20260822175138_17_13", role="在日历内浏览贴纸并拖放到日期；是情绪化层，但仍附着在时间结构上。", layout="托盘固定在日历底部、底栏上方；分类图标一行，贴纸网格在上；日历仍可见。", detail="贴纸本身可彩色，容器仍白底/轻线；分类当前态用下划线。", keys="分类=切换贴纸组；贴纸拖动=进入放置；+=商店/管理；时钟=历史/最近。", chain="月视图 → 贴纸键 → 托盘 → 选择/拖动贴纸 → 命中日期。", state="拖动时需要拾取反馈、目标日期高亮和无效区域回弹。", looka="情绪表达集中在资产，不扩散到底层组件。"),
    dict(section="E 贴纸", title="E2 贴纸落点 Popover：从装饰到日程", image="20260823124853_89_18", role="贴纸落到日期后，提供“登记日程”或“删除”等上下文动作。", layout="小型锚定 Popover 紧邻贴纸/日期；日历与托盘仍作为背景上下文。", detail="Popover 尺寸由短动作决定，不升级为全屏 Dialog；锚点方向避免遮住目标。", keys="登记日程=打开已带日期/贴纸的日程编辑器；删除=移除贴纸；点外部=关闭。", chain="拖放贴纸 → 日期命中 → Popover → 登记日程 → 新建日程（已预填） → 保存。", state="纯装饰贴纸与绑定日程贴纸需有可识别状态；删除语义不同。", looka="上下文动作用锚定 Popover，减少跳转。"),
    dict(section="E 贴纸", title="E3 贴纸绑定日程：保留来源信息", image="20260823124854_90_18", role="把贴纸的自由放置转成结构化日程，同时保留日期和资产。", layout="复用标准日程父编辑器；贴纸作为标题/装饰预填或关联字段，不另造页面。", detail="结构字段仍用白底表单；贴纸不成为大面积背景。", keys="保存=创建绑定日程；X=取消并回到日历贴纸上下文；字段行=子编辑器。", chain="贴纸 Popover → 登记日程 → 预填编辑器 → 保存 → 日历显示绑定状态。", state="取消时是否保留纯贴纸必须一致；建议保留已放置资产、只取消日程绑定。", looka="自由感与数据结构通过同一父编辑器衔接。"),
    dict(section="F 更多与设置", title="F1 更多：低频功能路由页", image="20260820133746_19_18", role="汇总账户、方案、设置、通知、贴纸/主题商店、帮助与消息。", layout="顶部黑色三项主入口；其余九宫格/分组入口；活动 Banner 位于内容区。", detail="黑色顶部建立明确层级，但不延伸到其他高频页面；通知角标只出现在有新内容的入口。", keys="账户/方案/设置=全页；商店=商品列表；帮助=说明页；消息=列表；底栏=返回高频域。", chain="底栏 More → 对应入口 → 子页 → X/← 返回 More。", state="角标计数在读后消失；外部链接需明确离开应用。", looka="低频入口集中，不污染底部导航。"),
    dict(section="F 更多与设置", title="F2 设置总览：Preference Router", image="20260820133745_18_18", role="按设置类型选择正确交互组件，而不是所有行都弹同一种 Dialog。", layout="全页列表按主题分组；行内开关、摘要行、普通路由行混合；X 返回 More。", detail="开关靠右；摘要位于主标签下；分组靠留白与小图标，不靠卡片。", keys="开关=即时写入；单选行=选择即关闭；多选=草稿+OK；复杂设置=全页；受限项=能力门。", chain="More → 设置 → 点偏好项 → 对应组件 → 提交/返回 → 设置页更新摘要。", state="按下态先出现，再打开 Dialog；开关失败原位回滚。", looka="建立设置组件路由表并全局复用。"),
    dict(section="F 更多与设置", title="F3 日历设置：长列表的分组密度", image="20260820133747_20_18", role="管理日历显示、周起始、节日、尺寸、默认值、同步与通知。", layout="长列表单列；每项可有摘要；类别用小图标/分组标题。", detail="行高48–56dp；主标签14–16sp，摘要12sp；开关与箭头不混用。", keys="摘要行=Picker/Dialog；开关=即时；复杂同步/通知=全页；X=返回设置总览。", chain="设置 → 日历基本设置 → 修改项 → 返回后行摘要同步。", state="外部同步开关需要加载/失败态；不可无反馈静默失败。", looka="长设置页依靠一致行模板和摘要，不依靠卡片。"),
    dict(section="F 更多与设置", title="F4 单选 Dialog：选择即关闭", image="20260822203452_31_4", role="月视图显示方式等少量互斥选项。", layout="标准86%宽、约18%高；radio 左置；当前项清楚。", detail="内容少时高度收敛；动作区可省略 OK。", keys="选项=提交并关闭；遮罩/返回=取消。", chain="设置行 → 单选 Dialog → 选择 → 关闭 → 行摘要更新。", state="Pressed 反馈先于 Overlay 首帧；焦点返回原行。", looka="互斥少量选项统一 select-and-close。"),
    dict(section="F 更多与设置", title="F5 视觉选择 Dialog：先看效果再确认", image="20260822203453_32_4", role="日期选择显示方式等需要视觉理解的选项。", layout="约45%屏高；两列预览图+单选；底部动作/说明。", detail="预览图承担主要解释，文字短；选中标记与图片边界清晰。", keys="预览=更新草稿；OK/确认=提交；取消=回滚。", chain="设置行 → 视觉选择 → 预览选择 → 确认 → 设置摘要/实际视图更新。", state="选择前后预览不应触发整页跳转；取消必须回旧值。", looka="视觉型偏好使用 preview + commit。"),
    dict(section="F 更多与设置", title="F6 文字大小 Dialog：语义级别而非任意滑杆", image="20260822203453_33_4", role="选择日程文字大小，选项为大/中/小。", layout="标准单选 Dialog；三项纵排；高度随内容收敛。", detail="当前值radio清晰；不使用过度精细的连续滑杆。", keys="大/中/小=提交并关闭；遮罩/返回=取消。", chain="日历设置 → 文字大小 → 选择 → 日历立即使用对应语义 Token。", state="预览或实际更新需即时可见；无障碍字体放大另行处理。", looka="尺寸映射到语义 Token，不存任意像素。"),
    dict(section="G 账户与内容", title="G1 账户：身份与同步能力入口", image="20260826112955_91_4", role="承载登录身份、同步和账户级能力；与本地日历设置分开。", layout="全页顶栏返回；身份信息在顶部；同步/方案作为后续入口。", detail="未登录态保持简洁，不用大面积营销卡；主动作清楚。", keys="返回=More；登录/注册=认证流程；同步=状态页；方案=订阅页。", chain="More → 账户 → 登录/同步/方案 → 完成 → 返回账户/More。", state="未登录、登录中、已登录、同步失败四态必须区分。", looka="账户能力是路由，不和普通设置混排。"),
    dict(section="G 账户与内容", title="G2 方案：价值说明与比较", image="20260820133750_23_18", role="说明基础/高级方案差异并进入购买或登录。", layout="插画/主张在上，比较信息在下；主动作突出但不做持续弹窗。", detail="颜色与品牌插画可比工具页更活泼；仍保持列对齐和有限按钮。", keys="方案比较=下移/展开；购买=支付；登录=账户；X=返回 More。", chain="More/能力门 → 方案 → 比较 → 购买/登录 → 成功 → 返回原意图。", state="价格、权益和恢复购买需以实时数据为准；成功后刷新能力。", looka="商业页可有情绪表达，但入口触发要克制。"),
    dict(section="G 账户与内容", title="G3 贴纸商店：搜索、分类、商品详情", image="20260820133755_27_18", role="浏览贴纸资产，按热门/新上架等分类进入商品详情。", layout="顶部搜索；分类页签；商品纵列/网格；价格与缩略图直接可见。", detail="商品图可丰富，壳仍白底；页签和搜索保持工具式。", keys="搜索=筛选；页签=切榜单；商品=详情；购买=系统支付；X=More。", chain="More → 贴纸商店 → 搜索/分类 → 商品 → 购买 → 下载 → 贴纸托盘可用。", state="购买中、已购、下载中、失败、恢复购买要分态。", looka="商店资产丰富，导航与交易反馈保持克制。"),
    dict(section="G 账户与内容", title="G4 主题活动：视觉预览与入口", image="20260820133749_22_18", role="主题促销/免费活动页，提供预览与选择。", layout="活动主视觉在上；颜色主题卡片/选项在下；关闭固定左上。", detail="此页允许高情绪色彩，但不把活动样式带入设置页。", keys="主题卡=预览/选择；领取/应用=提交；X=返回 More。", chain="More/主题 → 活动 → 选择 → 预览 → 应用 → 日历主题更新。", state="领取和应用分开；已拥有时不重复购买。", looka="营销视觉限制在内容页，工具骨架保持一致。"),
    dict(section="G 账户与内容", title="G5 主题应用结果：结构不因皮肤改变", image="20260826124641_114_4", role="展示主题被应用后的日历；皮肤改变情绪，不改变布局和交互位置。", layout="月份、网格、底栏位置与默认主题一致；装饰集中在顶/底或少量日期资产。", detail="文字对比度必须守住；装饰不遮挡日期/事项；当前态仍清晰。", keys="所有日历按键与默认主题一致；主题只改变视觉资产。", chain="主题详情 → 应用 → 返回月视图 → 原日期与事项仍可操作。", state="加载主题失败回退默认资源；不破坏数据可读性。", looka="主题系统只替换 Token/资产，不分叉页面结构。"),
    dict(section="G 账户与内容", title="G6 帮助：搜索优先的信息页", image="20260820133752_24_18", role="用搜索和分组文章解释应用使用方法。", layout="顶栏返回；搜索置顶；文章按主题分组的白底列表。", detail="内容型页面使用更宽松行距；链接蓝只用于可点击文本。", keys="搜索=帮助域检索；文章=详情/网页；返回=More。", chain="More → 帮助 → 搜索/主题 → 文章 → 返回搜索结果位置。", state="外链加载失败要可重试；离开应用需提示。", looka="帮助页以信息检索为主，不做卡片瀑布流。"),
    dict(section="G 账户与内容", title="G7 消息列表：新旧状态清楚", image="20260820133753_25_18", role="承载更新、运营与重要通知。", layout="顶栏返回；按时间倒序的标题+日期列表；新消息可用小标签/字重。", detail="不使用大红点覆盖正文；未读差异只需一个视觉变量。", keys="消息行=详情；返回=More并同步角标；可能的筛选保持次要。", chain="More（角标） → 消息列表 → 详情 → 返回 → 已读状态更新。", state="已读写入失败不影响阅读，但需在下次同步修复。", looka="未读状态克制、可扫读。"),
    dict(section="G 账户与内容", title="G8 消息详情：内容阅读优先", image="20260820133754_26_18", role="展示公告正文与必要插图/链接。", layout="标题、日期、正文纵向阅读；顶栏只保留返回。", detail="正文16sp左右、行距1.5附近；段落间距比设置页更宽。", keys="返回=消息列表；正文链接=外部/内部目标；分享如有则放更多。", chain="消息列表 → 消息详情 → 外链/返回 → 列表原位置。", state="长文加载与离线缓存；外链失败不丢当前阅读位置。", looka="内容详情减少操作 chrome。"),
    dict(section="G 账户与内容", title="G9 口令锁：系统级焦点与主题一致", image="20260826124642_115_4", role="在应用入口提供本地隐私保护。", layout="主题视觉可延续，但数字键盘与输入状态结构稳定。", detail="口令位数与错误反馈清楚；数字触点≥44dp；删除键位置固定。", keys="数字=输入；删除=退格；成功=进入应用；返回策略按安全要求。", chain="启动/恢复 → 口令页 → 输入正确 → 恢复上次页面；错误 → 原位反馈。", state="错误不清空过快；连续失败可限频；系统键盘不与自定义键盘叠加。", looka="安全页可换皮肤，不改变键位和反馈规则。"),
]


def add_screen_page(doc: Document, screen: dict, index: int) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(screen["section"])
    set_run_font(r, size=9, bold=True, color=CORAL)
    h = doc.add_heading(screen["title"], level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(7)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [3456, 5904], indent=120, border=True)
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(left, WHITE)
    set_cell_shading(right, "FBFCFD")
    set_cell_margins(left, top=110, start=120, bottom=100, end=120)
    set_cell_margins(right, top=120, start=170, bottom=110, end=170)

    img_path = find_image(screen["image"])
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(img_path), width=Inches(2.16))
    cp = left.add_paragraph(style="Evidence Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(f"图 {index:02d}｜B 级截图证据\n{img_path.name}")

    add_label_paragraph(right, "页面职责", screen["role"])
    add_label_paragraph(right, "布局比例", screen["layout"])
    add_label_paragraph(right, "UI 细节", screen["detail"])
    add_label_paragraph(right, "功能键逻辑", screen["keys"], color=DARK_BLUE)
    add_label_paragraph(right, "完整链路", screen["chain"], color=DARK_BLUE)
    add_label_paragraph(right, "状态与反馈", screen["state"])
    add_label_paragraph(right, "Looka 统一", screen["looka"], color=POSITIVE)

    p = doc.add_paragraph(style="Evidence Caption")
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("证据说明：页面结构与可见控件来自截图；动作结果结合配套录屏交叉确认。未能直接证明的时序/异常策略以 C 级建议表述。")


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("跨页面功能键总表", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1600, 2200, 2860, 2700])
    for i, h in enumerate(("功能键", "默认语义", "提交边界", "返回/失败")):
        set_cell_shading(table.cell(0, i), BLUE_GRAY)
        set_run_font(table.cell(0, i).paragraphs[0].add_run(h), size=9.2, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("X", "取消创建/关闭全页编辑器", "不提交父对象", "脏数据先确认；系统返回一致"),
        ("←", "返回上一级", "保留已明确保存的数据", "恢复来源筛选、滚动、选中日"),
        ("保存", "提交当前父对象", "整条日程/任务/笔记", "加载、防重复、失败留在原页"),
        ("完成/OK", "提交当前子编辑器或选择草稿", "仅当前字段/偏好", "回父页并更新摘要"),
        ("中央 +", "创建入口", "先选对象类型", "取消回当前页"),
        ("完成圆点", "切换任务完成", "即时写入", "失败原位回滚"),
        ("星标", "切换优先级", "即时写入", "状态和列表数量同步"),
        ("更多 …", "对象级低频动作", "菜单本身不提交", "点外部关闭；删除再确认"),
        ("开关", "布尔偏好", "即时写入", "失败回滚并给轻提示"),
        ("搜索", "当前功能域过滤", "输入即更新/防抖", "清空恢复原列表位置"),
        ("相机", "媒体附件/能力门", "权限成功后附加", "拒绝后解释并回编辑器"),
        ("贴纸", "打开托盘/拖放", "落点后成为装饰或绑定", "无效落点回弹"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_run_font(cells[i].paragraphs[0].add_run(text), size=8.7, color=BLACK)

    doc.add_page_break()
    doc.add_heading("完整关键链路", level=1)
    chains = [
        ("创建日程", "Calendar/中央 + → 日程创建 → 字段子编辑器 → 字段完成 → 父页保存 → 来源日历"),
        ("创建任务", "ToDo/中央 + → 任务创建 → 选清单/提醒/重复 → 父页保存 → 来源智能视图或清单"),
        ("建任务清单", "ToDo 首页/清单选择器 → 创建清单 → 名称 + 颜色 → 保存 → 持久清单并自动可选"),
        ("创建笔记列表", "Note 容器页/选择器 → 创建列表 → 输入 → 保存 → 持久列表 → 自动进入/选中"),
        ("贴纸登记日程", "Calendar → 贴纸托盘 → 拖放日期 → Popover → 登记日程 → 预填编辑器 → 保存"),
        ("删除对象", "详情/更多 → 删除 → 确认 Dialog → 提交 → 来源页 + Snackbar/撤销"),
        ("修改偏好", "More → 设置 → 偏好行 → Router 选择组件 → 提交 → 设置摘要/实际界面同步"),
        ("能力解锁", "受限动作 → 能力门 → 方案/账户 → 完成 → 返回并恢复原意图"),
    ]
    for label, text in chains:
        add_callout(doc, label, text, fill=CALLOUT)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    doc.add_page_break()
    doc.add_heading("全局验收清单", level=1)
    checks = [
        "同级页面的顶栏高度、左右安全边距、标题基线一致。",
        "同类 Dialog 的宽度约86%屏宽、低圆角、24–26dp内边距一致。",
        "Bottom Sheet、Popover、Dialog 的适用边界清楚，不互相代替。",
        "一个时刻只有一个交互 Overlay、一个焦点拥有者、一个 IME 拥有者。",
        "保存提交父对象；完成/OK只提交当前子编辑器；X与返回语义不混淆。",
        "高频状态（完成、星标、开关）即时；低频动作（复制、删除、排序）进入动作层。",
        "从详情/编辑返回时恢复来源页的筛选、滚动、清单或选中日期。",
        "空态、加载、错误、成功、删除中分别有独立反馈，不用同一灰态替代。",
        "点击热区不小于44×44dp；图标尺寸与热区分开定义。",
        "主题、贴纸、周末色只提供情绪/语义，不改变页面结构或降低文字对比。",
    ]
    for idx, item in enumerate(checks, 1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(f"{idx:02d}  ")
        set_run_font(r, size=10, bold=True, color=CORAL)
        set_run_font(p.add_run(item), size=10.5, color=BLACK)

    doc.add_heading("证据索引", level=2)
    p = doc.add_paragraph()
    set_run_font(p.add_run("截图："), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run("工作区“参考组件图标”目录中列示的 Lifebear 实机截图；每页图注保留原文件名。"), color=BLACK)
    p = doc.add_paragraph()
    set_run_font(p.add_run("录屏："), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run("设置 7f9d7827f156eb3d1e355221127db65e.mp4；笔记 0dffdedc81056c150141ff4a1153815e.mp4；任务 df9a1e8a886ad6b560e2fbd23e83dc94.mp4；日程/贴纸 ac3e6f1d560d921bf7b0ba0b4cf031b9.mp4。"), color=BLACK)
    p = doc.add_paragraph(style="Evidence Caption")
    p.add_run("注意：本报告是竞品研究与 Looka 统一设计依据，不是 Lifebear 官方设计规范。")


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    props = doc.core_properties
    props.title = "Lifebear UI/UX 逐页面图鉴与功能链路说明"
    props.subject = "Lifebear 页面级 UI/UX 研究；供 Looka 统一设计参考"
    props.author = "Codex"
    props.keywords = "Lifebear, UIUX, Looka, 页面图鉴, 功能链路"

    add_title_page(doc)
    add_front_matter(doc)
    for index, screen in enumerate(SCREENS, 1):
        add_screen_page(doc, screen, index)
    add_appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
